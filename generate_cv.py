from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── couleurs ──────────────────────────────────────────────
NAVY      = RGBColor(0x0D, 0x1B, 0x2A)
TEAL      = RGBColor(0x00, 0xC9, 0xA7)
TEAL_DARK = RGBColor(0x00, 0xA3, 0x8A)
GREY      = RGBColor(0x47, 0x55, 0x69)
LIGHT     = RGBColor(0xF8, 0xFA, 0xFC)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), val.get('sz', '4'))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=10,
            color=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def section_title(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = TEAL
    # bottom border via paragraph border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '00C9A7')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def bullet(cell, text, color=GREY, size=8.5):
    p = cell.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color

def normal_p(cell, text, size=8.5, color=GREY, space_before=2, space_after=2, bold=False, italic=False):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p

def clear_cell(cell):
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)


# ══════════════════════════════════════════════════════════
doc = Document()

# ── marges page ──────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Cm(0)
section.right_margin  = Cm(0)
section.top_margin    = Cm(0)
section.bottom_margin = Cm(0)

# Force A4 via XML (11906 x 16838 twips)
sectPr = section._sectPr
pgSz = sectPr.find(qn('w:pgSz'))
if pgSz is None:
    pgSz = OxmlElement('w:pgSz')
    sectPr.append(pgSz)
pgSz.set(qn('w:w'), '11906')
pgSz.set(qn('w:h'), '16838')

# ══════════════════════════════════════════════════════════
# TABLEAU PRINCIPAL  (2 colonnes : sidebar | contenu)
# ══════════════════════════════════════════════════════════
tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.allow_autofit = False

# Force la ligne à remplir toute la hauteur A4
tr = tbl.rows[0]._tr
trPr = tr.get_or_add_trPr()
trHeight = OxmlElement('w:trHeight')
trHeight.set(qn('w:val'), '16838')
trHeight.set(qn('w:hRule'), 'exact')
trPr.append(trHeight)

col_widths = [Cm(6.2), Cm(14.8)]
for i, cell in enumerate(tbl.rows[0].cells):
    cell.width = col_widths[i]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(col_widths[i].pt * 20 / 0.75)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

sidebar, main = tbl.rows[0].cells

# ══════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════
set_cell_bg(sidebar, '0D1B2A')
sidebar.width = col_widths[0]
for p in sidebar.paragraphs:
    p._element.getparent().remove(p._element)

sidebar._element.get_or_add_tcPr()
sidebar.vertical_alignment = WD_ALIGN_VERTICAL.TOP

def sb(text='', bold=False, italic=False, size=9.5, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=1, space_after=1):
    p = sidebar.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.3)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color or WHITE
    return p

def sb_title(text):
    p = sb(text.upper(), bold=True, size=8, color=TEAL, space_before=8, space_after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '00C9A7')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ─── Photo ──────────────────────────────────────
p = sidebar.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run()
run.add_picture('static/img/photo.jpg', width=Cm(4.5))

# ─── Nom & titre ────────────────────────────────
p = sb()
p.paragraph_format.space_before = Pt(2)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = sidebar.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(1)
run = p.add_run('Jesse Richard')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = WHITE

p = sidebar.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run('Développeur IA')
run.font.size = Pt(8)
run.font.color.rgb = TEAL

# ─── Contact ────────────────────────────────────
sb_title('Contact')
for label, value in [
    ('Tél.', '06 48 14 58 46'),
    ('Email', 'jessesteven26@gmail.com'),
    ('LinkedIn', 'jesse-richard'),
    ('Ville', 'Besançon (25) · Montcourt (70)'),
]:
    p = sidebar.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.4)
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = TEAL
    r2 = p.add_run(value)
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ─── Compétences techniques ──────────────────────
sb_title('Compétences')
for cat, skills in [
    ('ML',      'scikit-learn, pandas, numpy, clustering, régression'),
    ('DL',      'TensorFlow, PyTorch, Computer Vision'),
    ('Agents',  'LLM, systèmes agentiques, Stable Diffusion'),
    ('APIs',    'FastAPI, Flask, Django'),
    ('Data',    'SQL, PostgreSQL, Power BI, TABLEAU, Jupyter'),
    ('DevOps',  'Docker, Git, Make, CI/CD'),
    ('Langages','Python, SQL, HTML/CSS/JS'),
]:
    p = sidebar.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(cat + ' : ')
    r1.bold = True
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    r2 = p.add_run(skills)
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ─── Langues ────────────────────────────────────
sb_title('Langues')
for lang, level in [('Français', 'Natif'), ('Anglais', 'Technique (70%)')]:
    p = sidebar.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(lang + ' — ')
    r1.bold = True
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    r2 = p.add_run(level)
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = TEAL

# ─── Soft Skills ────────────────────────────────
sb_title('Soft Skills')
for sk in ['Autodidacte', 'Curieux', 'Sang-froid', 'Adaptable', 'Encadrement', 'Intelligence émotionnelle']:
    p = sidebar.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run('▸  ' + sk)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ─── Engagements ────────────────────────────────
sb_title('Engagements')
for eng in ['Bénévole associatif (JAL)', 'Accompagnateur centre de loisirs et scolaire', 'Scrutateur électoral']:
    p = sidebar.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run('▸  ' + eng)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ─── Citation ───────────────────────────────────
sb_title('Citation')
p = sidebar.add_paragraph()
p.paragraph_format.left_indent  = Cm(0.5)
p.paragraph_format.right_indent = Cm(0.3)
p.paragraph_format.space_before = Pt(2)
run = p.add_run('« Il faut être enthousiaste de son métier pour y exceller. »')
run.italic = True
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
p2 = sidebar.add_paragraph()
p2.paragraph_format.left_indent = Cm(0.5)
r2 = p2.add_run('— Denis Diderot')
r2.font.size = Pt(6.5)

# ─── QR code ────────────────────────────────────
p_qr = sidebar.add_paragraph()
p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_qr.paragraph_format.space_before = Pt(10)
p_qr.paragraph_format.space_after  = Pt(2)
run_qr = p_qr.add_run()
run_qr.add_picture('static/img/qr.png', width=Cm(3.5))

p_qr2 = sidebar.add_paragraph()
p_qr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_qr2.paragraph_format.space_before = Pt(0)
p_qr2.paragraph_format.space_after  = Pt(2)
r_qr2 = p_qr2.add_run('📱 Scanner → Portfolio interactif')
r_qr2.font.size = Pt(7.5)
r_qr2.italic = True
r_qr2.font.color.rgb = TEAL
r2.font.color.rgb = TEAL

# ══════════════════════════════════════════════════════════
# CONTENU PRINCIPAL
# ══════════════════════════════════════════════════════════
set_cell_bg(main, 'FFFFFF')
main.width = col_widths[1]
for p in main.paragraphs:
    p._element.getparent().remove(p._element)
main.vertical_alignment = WD_ALIGN_VERTICAL.TOP

def mp(text='', bold=False, italic=False, size=10, color=None,
       space_before=2, space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = main.add_paragraph()
    p.alignment = align
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color or GREY
    return p

def mp_title(text):
    p = mp(text.upper(), bold=True, size=8.5, color=NAVY, space_before=10, space_after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0D1B2A')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ─── En-tête ────────────────────────────────────
p = main.add_paragraph()
p.paragraph_format.left_indent  = Cm(0.6)
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run('Jesse Richard')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = NAVY

p = main.add_paragraph()
p.paragraph_format.left_indent  = Cm(0.6)
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(1)
run = p.add_run('Développeur en Intelligence Artificielle  ·  Data Scientist')
run.font.size = Pt(11)
run.font.color.rgb = TEAL

p = main.add_paragraph()
p.paragraph_format.left_indent  = Cm(0.6)
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run('Disponible pour alternance · Août 2026')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = TEAL_DARK

# ─── Profil ─────────────────────────────────────
mp_title('Profil')
mp(
    'Développeur IA autodidacte et curieux, passionné par les systèmes agentiques et les workflows '
    'code/nocode. À l\'aise sur toute la chaîne IA : collecte de données, entraînement de modèles, '
    'intégration API et déploiement. J\'apprends de nouvelles techniques chaque jour et j\'explore '
    'constamment de nouvelles architectures.',
    size=10, space_before=3, space_after=2
)

# ─── Expérience ─────────────────────────────────
mp_title('Expérience Professionnelle')

# Sophysa
p = main.add_paragraph()
p.paragraph_format.left_indent  = Cm(0.6)
p.paragraph_format.right_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(0)
r1 = p.add_run('Développeur en Intelligence Artificielle')
r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
r2 = p.add_run('   ·   ')
r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
r3 = p.add_run('Stage · 2025–2026')
r3.font.size = Pt(9); r3.font.color.rgb = TEAL_DARK

p2 = main.add_paragraph()
p2.paragraph_format.left_indent  = Cm(0.6)
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(2)
r = p2.add_run('Sophysa')
r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = TEAL_DARK

mp(
    'Développement de 11 algorithmes de clustering avec 4 types de métriques de distance '
    'pour l\'analyse de données. Conception d\'un frontend permettant l\'interaction et la '
    'visualisation des résultats de clustering. Déploiement containerisé avec Docker.',
    size=9, space_before=1, space_after=2
)
p = main.add_paragraph()
p.paragraph_format.left_indent  = Cm(0.6)
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after  = Pt(3)
for tag in ['scikit-learn', 'pandas', 'numpy', 'Clustering', 'Docker', 'Frontend data']:
    run = p.add_run(f'  {tag}  ')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

# CCD SNU
p = main.add_paragraph()
p.paragraph_format.left_indent  = Cm(0.6)
p.paragraph_format.right_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(0)
r1 = p.add_run('Cadre Encadrant SNU')
r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = NAVY
r2 = p.add_run('   ·   ')
r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
r3 = p.add_run('Mission civique')
r3.font.size = Pt(8); r3.font.color.rgb = TEAL_DARK

p2 = main.add_paragraph()
p2.paragraph_format.left_indent  = Cm(0.6)
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(2)
r = p2.add_run('CCD — Service National Universel · Vesoul, Lycée Belin')
r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = TEAL_DARK

mp(
    'Encadrement de jeunes volontaires au sein d\'un centre de cohésion. '
    'Gestion de section, organisation d\'activités, coordination pédagogique.',
    size=9, space_before=1, space_after=2
)

# ─── Formation ──────────────────────────────────
mp_title('Formation')

for titre, ecole, date, note in [
    (
        'GRETA — Développeur en Intelligence Artificielle',
        'GRETA Besançon × Simplon.co · Lycée Jules Haag — Titre RNCP Niveau 6',
        '2025–2027',
        '✓ Certification Microsoft Azure AI-900'
    ),
    (
        'Licence de Géologie (jusqu\'en L2)',
        'Université de Franche-Comté (UFC), Besançon',
        '2020–2025',
        ''
    ),
    (
        'Baccalauréat Scientifique — filière SVT',
        'Lycée Édouard Belin',
        '2020',
        ''
    ),
    (
        'BIA — Brevet d\'Initiation Aéronautique',
        '',
        '2016',
        ''
    ),
]:
    p = main.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(titre)
    r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(f'   {date}')
    r2.font.size = Pt(9); r2.font.color.rgb = TEAL_DARK

    if ecole:
        p2 = main.add_paragraph()
        p2.paragraph_format.left_indent  = Cm(0.6)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        r = p2.add_run(ecole)
        r.font.size = Pt(8.5); r.font.color.rgb = GREY

    if note:
        p3 = main.add_paragraph()
        p3.paragraph_format.left_indent  = Cm(0.6)
        p3.paragraph_format.space_before = Pt(1)
        p3.paragraph_format.space_after  = Pt(1)
        r = p3.add_run('▸ ' + note)
        r.italic = True; r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)

# ─── Projets ────────────────────────────────────
mp_title('Projets Réalisés')

projects = [
    ('🤖 IA Speaker — Agent LLM vocal',
     'Agent IA avec transcription vocale, synthèse vocale et clonage de voix. Chatbot LLM autonome.',
     'LLM · TTS/STT · Clonage vocal · Python'),
    ('🎨 Stable Diffusion App',
     'Application web de génération d\'images IA, multi-méthodes, scalable, haute qualité.',
     'Stable Diffusion · Deep Learning · Python'),
    ('🏥 Datathon — Myopie / Glaucome',
     'Algorithme de prédiction d\'évolution de la myopie par analyse d\'imagerie rétinienne.',
     'Computer Vision · PyTorch · Imagerie médicale'),
    ('🛡️ Détection de fraude bancaire',
     'App web + modèle de classification entraîné sur dataset réel, score de confiance.',
     'scikit-learn · Flask · Classification'),
    ('🚗 Prédiction prix voitures d\'occasion',
     'App full-stack + API REST : login, dashboard statistique (Power BI / TABLEAU), choix du modèle ML.',
     'FastAPI · pandas · Power BI · TABLEAU'),
    ('📥 YouTube Downloader',
     'Automatisation de téléchargement YouTube avec système de triggers et déploiement Docker.',
     'Python · Docker · Automation · Triggers'),
    ('🖥️ SAO HUD — Interface desktop IA',
     'Interface HUD desktop inspirée de SAO avec widgets CPU/GPU/RAM, explorateur, launcher. Auto-programmation par agents IA (Claude, Groq) : chaque jour un agent génère et propose de nouveaux widgets automatiquement.',
     'PyQt6 · psutil · Agent IA · Claude · Groq · Windows API'),
]

for name, desc, tech in projects:
    p = main.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(name)
    r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = NAVY

    p2 = main.add_paragraph()
    p2.paragraph_format.left_indent  = Cm(0.6)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run(desc)
    r2.font.size = Pt(8.5); r2.font.color.rgb = GREY

    p3 = main.add_paragraph()
    p3.paragraph_format.left_indent  = Cm(0.6)
    p3.paragraph_format.space_before = Pt(1)
    p3.paragraph_format.space_after  = Pt(0)
    r3 = p3.add_run(tech)
    r3.italic = True; r3.font.size = Pt(8); r3.font.color.rgb = TEAL_DARK

# ─── CTA portfolio ──────────────────────────────
p = main.add_paragraph()
p.paragraph_format.left_indent  = Cm(0)
p.paragraph_format.right_indent = Cm(0)
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(0)

# fond sombre via shading sur un tableau 1x1
cta_tbl = doc.add_table(rows=1, cols=1)

# on enlève ce paragraphe flottant et on l'insère dans la cellule main
# plus simple : paragraphe avec fond via XML
p2 = main.add_paragraph()
p2.paragraph_format.left_indent  = Cm(0.3)
p2.paragraph_format.right_indent = Cm(0.3)
p2.paragraph_format.space_before = Pt(8)
p2.paragraph_format.space_after  = Pt(6)
pPr2 = p2._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), '0D1B2A')
pPr2.append(shd)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

r1 = p2.add_run('👉  Version interactive du CV et portfolio :')
r1.bold = True
r1.font.size = Pt(8.5)
r1.font.color.rgb = WHITE

from docx.oxml import OxmlElement as _OE
br = _OE('w:br')
r1._r.append(br)

r2 = p2.add_run('cv-jesse-richard.onrender.com')
r2.bold = True
r2.font.size = Pt(9)
r2.font.color.rgb = TEAL
r2.underline = True

cta_tbl._element.getparent().remove(cta_tbl._element)


# ══════════════════════════════════════════════════════════
# SAUVEGARDE
# ══════════════════════════════════════════════════════════
output = 'CV_Jesse_Richard_Developpeur_IA.docx'
doc.save(output)
print(f'CV généré : {output}')
